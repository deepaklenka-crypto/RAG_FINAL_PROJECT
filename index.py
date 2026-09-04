"""
Document Ingestion & Indexing Pipeline:
Parses PDF, CSV, XLSX, DOCX, and TXT files.
Applies selectable chunking strategies (Recursive, Semantic, Structured).
Generates embeddings (Gemini text-embedding-004 or Local Fallback).
Upserts vectors and payloads into Qdrant Vector DB.
Populates PostgreSQL / SQLite metadata and Knowledge Graph tables.
"""

import os
import sys
import argparse
from typing import List, Dict, Any, Tuple
import pandas as pd
from pypdf import PdfReader
from docx import Document

from database import init_db, SessionLocal, DocumentModel, ChunkModel
from rag.chunking import get_chunker, TextChunk
from rag.embeddings import get_embedding_provider
from rag.vector_store import QdrantVectorStore
from rag.graph_rag import GraphRAG


class DocumentParser:
    """Extracts raw text and metadata from PDF, CSV, XLSX, DOCX, and TXT."""

    @staticmethod
    def parse(file_path: str) -> Tuple[str, Dict[str, Any]]:
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        metadata = {
            "filename": filename,
            "file_type": ext,
            "file_size": file_size,
            "file_path": file_path
        }

        if ext == "pdf":
            return DocumentParser._parse_pdf(file_path), metadata
        elif ext == "csv":
            return DocumentParser._parse_csv(file_path), metadata
        elif ext in ["xlsx", "xls"]:
            return DocumentParser._parse_xlsx(file_path), metadata
        elif ext in ["docx", "doc"]:
            return DocumentParser._parse_docx(file_path), metadata
        elif ext == "txt":
            return DocumentParser._parse_txt(file_path), metadata
        else:
            # Generic text fallback
            return DocumentParser._parse_txt(file_path), metadata

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(f"[Page {i+1}]\n{text.strip()}")
        return "\n\n".join(pages_text)

    @staticmethod
    def _parse_csv(file_path: str) -> str:
        df = pd.read_csv(file_path)
        lines = []
        for _, row in df.iterrows():
            record_str = ", ".join(f"{col}: {val}" for col, val in row.items())
            lines.append(record_str)
        return "\n".join(lines)

    @staticmethod
    def _parse_xlsx(file_path: str) -> str:
        excel = pd.ExcelFile(file_path)
        sections = []
        for sheet_name in excel.sheet_names:
            df = pd.read_excel(excel, sheet_name=sheet_name)
            sections.append(f"### Sheet: {sheet_name}")
            for _, row in df.iterrows():
                record_str = ", ".join(f"{col}: {val}" for col, val in row.items())
                sections.append(record_str)
        return "\n".join(sections)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        doc = Document(file_path)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for t in doc.tables:
            for row in t.rows:
                row_vals = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_vals:
                    lines.append(" | ".join(row_vals))
        return "\n\n".join(lines)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()


def index_document(
    file_path: str,
    chunking_strategy: str = "recursive",
    embedding_strategy: str = "gemini",
    chunk_size: int = 500,
    chunk_overlap: int = 100,
    extract_graph: bool = True
) -> Dict[str, Any]:
    """
    Main ingestion worker:
    1. Parses file content into text.
    2. Records document in PostgreSQL/SQLite.
    3. Splits text using selected chunker.
    4. Computes vector embeddings.
    5. Upserts chunks and vectors into Qdrant.
    6. Extracts Graph entities and relationships for Graph RAG.
    """
    init_db()

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # 1. Parse File
    text, meta = DocumentParser.parse(file_path)
    if not text.strip():
        return {"status": "skipped", "reason": "Empty document"}

    # 2. Record in DB
    with SessionLocal() as db:
        doc_record = DocumentModel(
            filename=meta["filename"],
            file_type=meta["file_type"],
            file_size=meta["file_size"],
            file_path=meta["file_path"],
            status="indexing",
            chunking_strategy=chunking_strategy,
            embedding_strategy=embedding_strategy
        )
        db.add(doc_record)
        db.commit()
        db.refresh(doc_record)
        doc_id = doc_record.id

    # 3. Chunking
    chunker = get_chunker(chunking_strategy, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks: List[TextChunk] = chunker.chunk(text, metadata={"document_id": doc_id, **meta})

    chunk_texts = [c.text for c in chunks]
    chunk_metas = [c.metadata for c in chunks]

    # 4. Embeddings
    embed_provider = get_embedding_provider(embedding_strategy)
    embeddings = embed_provider.embed_batch(chunk_texts)

    # 5. Qdrant Vector Store
    vector_store = QdrantVectorStore(dimension=embed_provider.dimension)
    assigned_ids = vector_store.upsert_chunks(
        chunks=chunk_texts,
        embeddings=embeddings,
        metadatas=chunk_metas
    )

    # 6. Save Chunk Records to DB
    with SessionLocal() as db:
        for i, chunk in enumerate(chunks):
            embedding_id = assigned_ids[i] if i < len(assigned_ids) else str(i)
            chunk_rec = ChunkModel(
                document_id=doc_id,
                chunk_index=chunk.chunk_index,
                content=chunk.text,
                token_count=chunk.token_count,
                embedding_id=embedding_id
            )
            db.add(chunk_rec)

        # Update doc status
        doc_update = db.query(DocumentModel).filter_by(id=doc_id).first()
        if doc_update:
            doc_update.status = "indexed"
            doc_update.chunk_count = len(chunks)
        db.commit()

    # 7. Graph RAG entity & relation extraction
    graph_stats = {"entities": 0, "relations": 0}
    if extract_graph:
        graph_rag = GraphRAG(vector_store=vector_store)
        for chunk in chunks[:8]:  # Index top chunks to knowledge graph
            stats = graph_rag.extract_and_index_graph(chunk.text, doc_id=doc_id)
            graph_stats["entities"] += stats["entities"]
            graph_stats["relations"] += stats["relations"]

    return {
        "status": "success",
        "document_id": doc_id,
        "filename": meta["filename"],
        "file_type": meta["file_type"],
        "chunks_indexed": len(chunks),
        "chunking_strategy": chunking_strategy,
        "embedding_strategy": embedding_strategy,
        "graph_stats": graph_stats
    }


def index_all_samples(samples_dir: str = "./data/samples"):
    """Convenience helper to index all sample files."""
    if not os.path.exists(samples_dir):
        print(f"Directory {samples_dir} does not exist.")
        return

    files = [os.path.join(samples_dir, f) for f in os.listdir(samples_dir)]
    print(f"Discovered {len(files)} sample files in {samples_dir}:")
    for f in files:
        print(f" - Indexing: {f}")
        try:
            res = index_document(f, chunking_strategy="recursive", extract_graph=True)
            print(f"   Indexed doc_id={res['document_id']}, chunks={res['chunks_indexed']}, graph={res['graph_stats']}")
        except Exception as e:
            print(f"   Failed to index {f}: {e}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Document Ingestion & Indexing Pipeline")
    parser.add_argument("--file", type=str, help="Path to single file to index")
    parser.add_argument("--all-samples", action="store_true", help="Index all test files in data/samples")
    parser.add_argument("--chunking", type=str, default="recursive", choices=["recursive", "semantic", "structured"])
    parser.add_argument("--embedding", type=str, default="gemini", choices=["gemini", "local"])

    args = parser.parse_args()

    if args.all_samples:
        index_all_samples()
    elif args.file:
        res = index_document(args.file, chunking_strategy=args.chunking, embedding_strategy=args.embedding)
        print("Indexing completed successfully:")
        print(res)
    else:
        parser.print_help()
