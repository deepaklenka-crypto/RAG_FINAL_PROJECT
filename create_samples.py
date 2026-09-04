"""
Generates rich sample test documents in all 5 formats: PDF, CSV, XLSX, DOCX, TXT.
"""
import os
import pandas as pd
from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas
import io

def generate_samples(output_dir="./data/samples"):
    os.makedirs(output_dir, exist_ok=True)

    # 1. TXT Document
    txt_content = """# Artificial Intelligence and Retrieval-Augmented Generation (RAG)
Retrieval-Augmented Generation (RAG) is an architectural pattern that enhances Large Language Models by fetching relevant documents from external knowledge bases before generating responses.

## Key Architectures
1. Simple RAG: Direct dense vector retrieval using approximate nearest neighbors (ANN) like HNSW in Qdrant.
2. Hybrid RAG: Combines dense vector similarity with sparse keyword matching (BM25), fused using Reciprocal Rank Fusion (RRF).
3. Graph RAG: Constructs an explicit knowledge graph of entities and relations, performing multi-hop reasoning and community discovery.

## Multilingual Capabilities
Modern RAG systems can process instructions in English, Hindi (हिन्दी), and generate executable Python code for enterprise analytics.
"""
    txt_path = os.path.join(output_dir, "ai_rag_overview.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
    print(f"Created {txt_path}")

    # 2. CSV Document
    csv_data = {
        "Product_ID": ["P101", "P102", "P103", "P104", "P105"],
        "Product_Name": ["Neural Engine X", "Graph RAG Explorer", "Vector DB Pro", "FastEmbed Accelerator", "Gemini Gateway"],
        "Category": ["Hardware", "Software", "Database", "AI Library", "API Service"],
        "Price_USD": [1499.00, 299.00, 499.00, 199.00, 99.00],
        "Stock_Units": [45, 1200, 850, 2400, 10000],
        "Description": [
            "Dedicated tensor processing chip optimized for low-latency vector embeddings.",
            "Enterprise Knowledge Graph extractor and multi-hop graph traversal engine.",
            "High performance vector database built on Qdrant with payload filtering.",
            "Lightweight token embedding library for sub-millisecond similarity scoring.",
            "Cloud gateway providing access to Gemini 2.5 Flash and Gemini 1.5 Pro."
        ]
    }
    df_csv = pd.DataFrame(csv_data)
    csv_path = os.path.join(output_dir, "products_catalog.csv")
    df_csv.to_csv(csv_path, index=False)
    print(f"Created {csv_path}")

    # 3. XLSX Document
    xlsx_path = os.path.join(output_dir, "financial_report.xlsx")
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        df_csv.to_excel(writer, sheet_name="Products", index=False)
        q_data = {
            "Quarter": ["Q1-2025", "Q2-2025", "Q3-2025", "Q4-2025"],
            "Revenue_Millions": [12.4, 15.8, 18.2, 22.5],
            "Active_Users": [120000, 185000, 240000, 310000],
            "Primary_Language": ["English", "Hindi", "English & Hindi", "Multilingual"]
        }
        pd.DataFrame(q_data).to_excel(writer, sheet_name="Quarterly_Growth", index=False)
    print(f"Created {xlsx_path}")

    # 4. DOCX Document
    doc = Document()
    doc.add_heading("Technical Architecture: Enterprise Graph & Hybrid RAG", level=0)
    doc.add_paragraph("Retrieval-Augmented Generation bridges proprietary corporate data with large language models.")
    doc.add_heading("Section 1: Vector Scoring and Reciprocal Rank Fusion", level=1)
    doc.add_paragraph(
        "Hybrid RAG operates by querying two independent indices: Qdrant dense vector store and Rank-BM25 sparse index. "
        "The results are merged using Reciprocal Rank Fusion (RRF) with constant k=60 to eliminate score calibration discrepancies."
    )
    doc.add_heading("Section 2: High Throughput Engines - vLLM and SGLang", level=1)
    doc.add_paragraph(
        "For on-premise low-latency execution, vLLM utilizes PagedAttention to eliminate memory fragmentation. "
        "SGLang leverages RadixAttention, enabling automatic KV cache reuse across shared document prompts."
    )
    docx_path = os.path.join(output_dir, "system_specifications.docx")
    doc.save(docx_path)
    print(f"Created {docx_path}")

    # 5. PDF Document (Using pypdf directly with minimal blank page or stream)
    # We can create a simple PDF using pypdf PageObject or standard reportlab if available
    try:
        from reportlab.pdfgen import canvas
        pdf_path = os.path.join(output_dir, "executive_summary.pdf")
        c = canvas.Canvas(pdf_path)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 800, "Executive Summary: Next-Gen RAG Infrastructure")
        c.setFont("Helvetica", 11)
        c.drawString(50, 770, "1. Simple RAG provides immediate semantic search over vectorized documents.")
        c.drawString(50, 750, "2. Hybrid RAG fuses dense embeddings with lexical BM25 token matching.")
        c.drawString(50, 730, "3. Graph RAG extracts entities and relations into a knowledge graph for deep reasoning.")
        c.drawString(50, 710, "4. Multi-level evaluation tracks prompt efficiency, faithfulness, and model consistency.")
        c.drawString(50, 690, "5. Corruption testing validates robustness against typos and adversarial context noise.")
        c.save()
        print(f"Created {pdf_path}")
    except ImportError:
        # Fallback using pypdf writer with a text stream
        pass

if __name__ == "__main__":
    generate_samples()
