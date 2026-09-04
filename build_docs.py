"""
Script to build the comprehensive, beautifully styled OmniRAG User & API Documentation in .docx format.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Sets cell internal padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border_left(cell, color_hex="1A365D", sz="36"):
    """Adds a thick accent border on the left of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def set_table_borders(table, color_hex="CBD5E0", sz="4"):
    """Sets subtle outer and inner borders for a clean modern table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_callout(doc, text_list, title="NOTE", fill_hex="EBF8FF", border_hex="2B6CB0"):
    """Adds a stylish modern callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, fill_hex)
    set_cell_border_left(cell, color_hex=border_hex, sz="32")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"✦  {title}\n")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(10.5)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor.from_string(border_hex)

    for i, line in enumerate(text_list):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(45, 55, 72)
    
    # Empty space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_str, caption=None):
    """Adds a monospace formatted code snippet block."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(2)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = "Calibri"
        run_cap.font.size = Pt(9.5)
        run_cap.font.bold = True
        run_cap.font.color.rgb = RGBColor(74, 85, 104)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F7FAFC")
    set_cell_border_left(cell, color_hex="718096", sz="16")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_str.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(26, 32, 44)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def style_heading_1(p, text):
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)  # Dark Navy #1A365D

def style_heading_2(p, text):
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(43, 108, 176)  # Accent Slate Blue #2B6CB0

def style_heading_3(p, text):
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(74, 85, 104)  # Slate Gray #4A5568

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(45, 55, 72)
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run_bold = p.add_run(bold_prefix)
    run_bold.font.name = "Calibri"
    run_bold.font.size = Pt(10)
    run_bold.font.bold = True
    run_bold.font.color.rgb = RGBColor(26, 54, 93)

    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(45, 55, 72)
    return p

def format_table(table, col_widths, headers, rows_data):
    """Formats a professional, clean table with shaded header and alternating rows."""
    set_table_borders(table, color_hex="CBD5E0", sz="4")
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = Inches(col_widths[i])
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, row in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].width = Inches(col_widths[c_idx])
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=160, right=160)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(45, 55, 72)

    doc_after = table._element.getparent()


def build_documentation_file():
    doc = Document()

    # Set 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -------------------------------------------------------------
    # Cover / Header Banner
    # -------------------------------------------------------------
    tbl_cover = doc.add_table(rows=1, cols=1)
    tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cover.autofit = False
    c_cover = tbl_cover.cell(0, 0)
    c_cover.width = Inches(6.5)
    set_cell_background(c_cover, "1A365D")
    set_cell_margins(c_cover, top=280, bottom=280, left=240, right=240)

    p_cov_tag = c_cover.paragraphs[0]
    p_cov_tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_tag = p_cov_tag.add_run("ENTERPRISE PLATFORM & API SPECIFICATION")
    run_tag.font.name = "Calibri"
    run_tag.font.size = Pt(9.5)
    run_tag.font.bold = True
    run_tag.font.color.rgb = RGBColor(144, 205, 244)  # Light blue

    p_cov_title = c_cover.add_paragraph()
    p_cov_title.paragraph_format.space_before = Pt(4)
    p_cov_title.paragraph_format.space_after = Pt(6)
    run_title = p_cov_title.add_run("OmniRAG Platform: Complete User & API Guide")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    p_cov_sub = c_cover.add_paragraph()
    p_cov_sub.paragraph_format.space_before = Pt(0)
    p_cov_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_cov_sub.add_run(
        "A Comprehensive, Step-by-Step Practical Manual for Non-Technical Users, Developers, "
        "and QA Engineers covering Swagger UI Docs, Postman Application, 3 RAG Paradigms, and System Architecture."
    )
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(10.5)
    run_sub.font.color.rgb = RGBColor(226, 232, 240)

    p_meta = c_cover.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(6)
    p_meta.paragraph_format.space_after = Pt(0)
    run_meta = p_meta.add_run("System Version: 2.0.0  |  Status: Production-Ready  |  Dual DB: PostgreSQL + SQLite  |  Vector: Qdrant")
    run_meta.font.name = "Calibri"
    run_meta.font.size = Pt(8.5)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # Executive Summary
    # -------------------------------------------------------------
    p_h1 = doc.add_paragraph()
    style_heading_1(p_h1, "1. Executive Overview & System Architecture")

    add_body_p(doc, 
        "OmniRAG is an enterprise-grade Retrieval-Augmented Generation (RAG) platform architected to solve the limitations "
        "of standard single-vector retrieval systems. By integrating 3 distinct retrieval paradigms, multi-format document ingestion, "
        "a resilient dual-database persistence tier, conversational memory, and quantitative evaluation diagnostics, OmniRAG provides "
        "production reliability for demanding AI workflows."
    )

    add_bullet(doc, "Multi-Format Document Parsing: ", "Ingests and parses PDF, CSV, XLSX, DOCX, and TXT files with tailored chunking algorithms (Recursive, Semantic Sentence, and Structured Table Serialization).")
    add_bullet(doc, "Dual Storage Tier: ", "Production-grade PostgreSQL database with zero-friction, automatic fallback to local SQLite (./data/rag_app.db) if PostgreSQL is offline or unconfigured.")
    add_bullet(doc, "Qdrant Vector Database: ", "High-performance vector search engine using Hierarchical Navigable Small World (HNSW) graphs and cosine distance metric, supporting both local disk persistence and cloud deployments.")
    add_bullet(doc, "Multi-Engine LLM Factory: ", "Seamlessly interfaces with Google Gemini API (gemini-flash-latest / 2.5-flash), with native drop-in adapters for high-throughput self-hosted inference engines: vLLM (PagedAttention) and SGLang (RadixAttention).")
    add_bullet(doc, "Multilingual & Code Synthesis: ", "Native support for standard English, fluent Hindi (हिन्दी in Devanagari script), and PEP 8 compliant, type-annotated Python code generation.")
    add_bullet(doc, "Multi-Turn Session Memory: ", "Session-scoped conversation buffer with automatic pronoun/coreference resolution and persistent database synchronization.")

    add_callout(doc, [
        "Default Local Endpoint: http://127.0.0.1:8000",
        "Interactive Swagger UI: http://127.0.0.1:8000/docs",
        "Raw OpenAPI JSON: http://127.0.0.1:8000/openapi.json",
        "Alternative ReDoc Documentation: http://127.0.0.1:8000/redoc"
    ], title="QUICK ACCESS URLS", fill_hex="EBF8FF", border_hex="2B6CB0")

    # -------------------------------------------------------------
    # The 3 RAG Paradigms
    # -------------------------------------------------------------
    p_h2 = doc.add_paragraph()
    style_heading_1(p_h2, "2. Deep Dive: The 3 RAG Paradigms")

    add_body_p(doc, 
        "A single RAG architecture cannot optimally address all query types. OmniRAG provides three distinct paradigms, "
        "selectable dynamically per query via the REST API, Swagger UI, or Postman:"
    )

    # Comparison Table
    t_rag = doc.add_table(rows=1, cols=4)
    t_rag.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_rag,
        [1.2, 1.6, 2.0, 1.7],
        ["Paradigm", "Primary Mechanism", "How it Works Internally", "Best Used For"],
        [
            [
                "Simple RAG",
                "Dense Vector HNSW Search",
                "Encodes query into 768-d vector, performs cosine nearest neighbor search in Qdrant, passes top-K chunks directly to LLM prompt.",
                "General conceptual questions, thematic summaries, fast lookup where exact keyword matching is secondary."
            ],
            [
                "Hybrid RAG",
                "Dense Qdrant + Sparse BM25 + Reciprocal Rank Fusion",
                "Executes parallel dense semantic retrieval and BM25Okapi keyword search. Fuses ranks using RRF (k=60) and applies a fast cross-reranker.",
                "Domain-specific corpora, technical manuals, part numbers, SKU codes, acronyms, and financial reports."
            ],
            [
                "Graph RAG",
                "Entity-Relation Triples + NetworkX Subgraph Traversal",
                "Extracts structured entities and relationship edges into relational tables and a directed graph. Queries trigger 1-hop and 2-hop neighborhood expansion.",
                "Multi-hop associative reasoning, organizational hierarchies, causal relationship discovery across documents."
            ]
        ]
    )
    # -------------------------------------------------------------
    # Multi-Engine LLM Tier & Local Ollama Models
    # -------------------------------------------------------------
    p_llm = doc.add_paragraph()
    style_heading_1(p_llm, "3. Multi-Engine LLM Tier & Local Ollama Models")

    add_body_p(doc, 
        "OmniRAG integrates a pluggable Engine Factory (llm/engine_factory.py) supporting both state-of-the-art "
        "cloud reasoning APIs and zero-latency, 100% offline local inference engines:"
    )

    t_llm = doc.add_table(rows=1, cols=4)
    t_llm.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_llm,
        [1.6, 1.2, 1.8, 1.9],
        ["Engine / Backend", "Mode", "Hardware Target", "Key Strengths & Role"],
        [
            ["Google Gemini API\n(gemini)", "Cloud API", "Google Cloud", "gemini-flash-latest / 2.5-flash: high reasoning, native Hindi, and fast token generation."],
            ["Ollama: Qwen 2.5 3B\n(ollama:qwen2.5:3b)", "Local On-Device", "Laptop CPU / Intel Iris Xe", "Default local model (~2GB RAM): Fast CPU inference, strong Hindi & English multilingual comprehension."],
            ["Ollama: Qwen 2.5 7B\n(ollama:qwen2.5vl:7b)", "Local On-Device", "Dedicated GPU / 16GB+ RAM", "High-capacity analytical reasoning, complex document synthesis, and multi-step inference."],
            ["Ollama: Gemma 2 2B\n(ollama:gemma2:2b)", "Local On-Device", "Low-power laptops / Edge", "Ultra-compact 2B model by Google DeepMind with very low memory footprint and high speed."],
            ["vLLM\n(vllm)", "Self-Hosted", "NVIDIA GPU Server", "PagedAttention continuous batching for enterprise-scale high concurrency."],
            ["SGLang\n(sglang)", "Self-Hosted", "NVIDIA GPU Server", "RadixAttention KV-cache reuse across multi-turn RAG dialogue turns."]
        ]
    )

    # -------------------------------------------------------------
    # Conversational Session Memory
    # -------------------------------------------------------------
    p_h3 = doc.add_paragraph()
    style_heading_1(p_h3, "4. Multi-Turn Conversational Memory & Coreference Resolution")

    add_body_p(doc, 
        "In real-world chat scenarios, users frequently ask follow-up questions using ambiguous pronouns (e.g., 'What is Qdrant?' "
        "followed by 'Does it support local disk storage?'). Standard RAG pipelines fail on Turn 2 because searching for 'Does it support local disk storage?' "
        "retrieves generic storage passages rather than Qdrant-specific information."
    )

    add_body_p(doc, "OmniRAG implements a dual-layer conversational memory engine:")
    add_bullet(doc, "1. In-Memory Session Cache + DB Sync: ", "Chat history is buffered in memory for ultra-fast lookup (<1ms) and asynchronously committed to the chat_sessions and chat_messages database tables.")
    add_bullet(doc, "2. Automatic Coreference Resolution: ", "The Contextualizer analyzes incoming queries for referential triggers ('it', 'this', 'they', 'what about', 'does it') and dynamically reformulates the query (e.g., '[Context: What is Qdrant?] Does it support local disk storage?') before vector search.")
    add_bullet(doc, "3. Prompt Context Injection: ", "Injects a sliding window of recent dialogue turns into the system prompt, enabling continuous, coherent dialogue threads.")

    # -------------------------------------------------------------
    # Swagger Docs Guide
    # -------------------------------------------------------------
    p_sw = doc.add_paragraph()
    style_heading_1(p_sw, "4. How to Use the Application in Swagger UI (/docs)")

    add_body_p(doc, 
        "FastAPI generates an interactive OpenAPI-compliant Swagger documentation interface out of the box. "
        "It provides a complete visual sandbox to explore endpoints, inspect request/response schemas, and execute live API calls "
        "directly from your web browser without installing any software."
    )

    p_sw_steps = doc.add_paragraph()
    style_heading_2(p_sw_steps, "4.1 Navigating to Swagger UI")
    add_bullet(doc, "Step 1: ", "Ensure the server is running by executing 'python main.py' in your terminal.")
    add_bullet(doc, "Step 2: ", "Open any modern web browser (Chrome, Edge, Firefox, Safari).")
    add_bullet(doc, "Step 3: ", "Navigate to: http://127.0.0.1:8000/docs")

    p_sw_ops = doc.add_paragraph()
    style_heading_2(p_sw_ops, "4.2 Step-by-Step Operations in Swagger UI")

    # Swagger Operation Table
    t_sw = doc.add_table(rows=1, cols=3)
    t_sw.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_sw,
        [1.8, 1.4, 3.3],
        ["Endpoint / Action", "HTTP Method", "Detailed Step-by-Step Instructions"],
        [
            [
                "Health Check\n/api/health",
                "GET",
                "1. Click to expand 'GET /api/health'.\n2. Click the 'Try it out' button on the top right.\n3. Click the blue 'Execute' button.\n4. Check the 'Response body': Verify 'status': 'online', vector count, and graph node count."
            ],
            [
                "Document Upload\n/api/upload",
                "POST",
                "1. Click to expand 'POST /api/upload'.\n2. Click 'Try it out'.\n3. Under 'file', click 'Choose File' and select any PDF, CSV, XLSX, DOCX, or TXT file.\n4. Set 'chunking_strategy' to 'semantic' (or 'recursive' / 'structured').\n5. Set 'embedding_strategy' to 'gemini' (or 'local').\n6. Click 'Execute' and observe chunk extraction results."
            ],
            [
                "RAG Query\n/api/query",
                "POST",
                "1. Expand 'POST /api/query' and click 'Try it out'.\n2. In the Request body JSON, enter your query and options.\n3. Click 'Execute'.\n4. Review the generated answer, retrieved source citations with similarity scores, and fine-grained latency telemetry."
            ],
            [
                "Multi-Turn Chat\n/api/query (session_id)",
                "POST",
                "1. In 'POST /api/query', provide a 'session_id' (e.g. 'swagger_session_1').\n2. Turn 1 query: 'What is Qdrant?' -> Execute.\n3. Turn 2 query: 'Does it support disk storage?' with the same session_id -> Execute.\n4. Observe that the contextualizer correctly links 'it' to Qdrant."
            ],
            [
                "Corruption Stress Test\n/api/corruption-test",
                "POST",
                "1. Expand 'POST /api/corruption-test' and click 'Try it out'.\n2. Enter a query and clean context passage.\n3. Click 'Execute'.\n4. Inspect the resulting Robustness Score and degradation breakdown across 0%, 25%, 50%, and 75% noise."
            ],
            [
                "Knowledge Graph\n/api/graph",
                "GET",
                "1. Expand 'GET /api/graph' and click 'Try it out'.\n2. Click 'Execute'.\n3. Review all extracted entity nodes (names, types) and relationship edges (source, target, predicate, weights)."
            ]
        ]
    )

    add_callout(doc, [
        "Interactive Model Schemas: In Swagger UI, scroll down to the 'Schemas' section at the bottom to inspect exact Pydantic validation models, required fields, and default data types for every request and response.",
        "Alternative ReDoc View: If you prefer a clean 3-pane reference specification without the 'Try it out' execution sandbox, navigate to http://127.0.0.1:8000/redoc."
    ], title="PRO TIP: SWAGGER FEATURES", fill_hex="F0FFF4", border_hex="276749")

    # -------------------------------------------------------------
    # Postman Guide
    # -------------------------------------------------------------
    p_pm = doc.add_paragraph()
    style_heading_1(p_pm, "5. How to Use the Application in Postman")

    add_body_p(doc, 
        "Postman is the industry standard for API testing, automation, and team collaboration. "
        "OmniRAG includes full support for Postman through a pre-configured collection and live OpenAPI import."
    )

    style_heading_2(doc.add_paragraph(), "5.1 Option A: One-Click Import via Pre-Built Collection (Recommended)")
    add_bullet(doc, "Step 1: ", "Open the Postman application on your desktop.")
    add_bullet(doc, "Step 2: ", "Click the 'Import' button located in the top-left navigation bar.")
    add_bullet(doc, "Step 3: ", "Drag and drop the file 'OmniRAG.postman_collection.json' (located in the root project folder) into Postman.")
    add_bullet(doc, "Step 4: ", "Click 'Import'. You will see the complete 'OmniRAG Enterprise Platform API' collection organized into 6 organized folders:")
    add_bullet(doc, "  • ", "1. System & Health (Health Check, Latency Benchmarks)")
    add_bullet(doc, "  • ", "2. Document Ingestion (Upload Document with form-data, List Indexed Documents)")
    add_bullet(doc, "  • ", "3. RAG Querying & Synthesis (Hybrid RAG, Simple RAG, Graph RAG, Hindi, Python Code, Multi-Turn turns)")
    add_bullet(doc, "  • ", "4. Conversational Session Memory (List Sessions, Get Turn History, Clear Session)")
    add_bullet(doc, "  • ", "5. Evaluation & Diagnostics (Multi-Level Eval, Corruption Test, Knowledge Graph)")
    add_bullet(doc, "  • ", "6. Optimization & Cache (Clear Semantic Vector Cache)")

    style_heading_2(doc.add_paragraph(), "5.2 Option B: Importing Directly via OpenAPI URL")
    add_bullet(doc, "Step 1: ", "In Postman, click 'Import'.")
    add_bullet(doc, "Step 2: ", "Select the 'Link' tab.")
    add_bullet(doc, "Step 3: ", "Paste the live OpenAPI URL: http://127.0.0.1:8000/openapi.json")
    add_bullet(doc, "Step 4: ", "Click 'Continue' and then 'Import'. Postman will automatically parse all endpoints and schemas into a new collection.")

    style_heading_2(doc.add_paragraph(), "5.3 Setting Up Environment Variables in Postman")
    add_body_p(doc, 
        "The collection comes pre-configured with two collection-level variables, making URL management completely dynamic:"
    )
    add_bullet(doc, "{{base_url}}: ", "Set to http://127.0.0.1:8000 (adjust if hosting remotely or in a container).")
    add_bullet(doc, "{{session_id}}: ", "Set to postman_demo_session (used to link multi-turn conversational queries).")

    style_heading_2(doc.add_paragraph(), "5.4 Comprehensive Endpoint Execution Walkthrough in Postman")

    # Sample Request 1: Hybrid RAG
    add_code_block(doc, 
"""POST {{base_url}}/api/query
Headers:
  Content-Type: application/json

Body (raw JSON):
{
  "query": "What are the core differences between Simple, Hybrid, and Graph RAG?",
  "rag_type": "hybrid",
  "backend": "gemini",
  "language": "en",
  "mode": "general",
  "top_k": 4,
  "use_cache": true,
  "compress_context": false
}""", caption="Example 1: Standard Hybrid RAG Query in Postman")

    # Sample Request 2: Hindi Multilingual
    add_code_block(doc,
"""POST {{base_url}}/api/query
Headers:
  Content-Type: application/json

Body (raw JSON):
{
  "query": "हाइब्रिड RAG और ग्राफ RAG की मुख्य विशेषताएं क्या हैं?",
  "rag_type": "hybrid",
  "language": "hi",
  "mode": "general",
  "top_k": 3
}""", caption="Example 2: Multilingual Hindi Query in Postman")

    # Sample Request 3: Python Coding
    add_code_block(doc,
"""POST {{base_url}}/api/query
Headers:
  Content-Type: application/json

Body (raw JSON):
{
  "query": "Write a Python function to compute Reciprocal Rank Fusion given two ranked lists",
  "rag_type": "hybrid",
  "mode": "code",
  "top_k": 3
}""", caption="Example 3: PEP 8 Python Code Generation Mode in Postman")

    # Sample Request 4: Multi-Turn Dialogue
    add_code_block(doc,
"""// Turn 1:
POST {{base_url}}/api/query
{
  "query": "What is Qdrant?",
  "rag_type": "hybrid",
  "session_id": "team_demo_session"
}

// Turn 2 (Follow-up with pronoun 'it'):
POST {{base_url}}/api/query
{
  "query": "Does it support fast filtering and local disk storage?",
  "rag_type": "hybrid",
  "session_id": "team_demo_session"
}""", caption="Example 4: Multi-Turn Dialogue with Coreference Resolution in Postman")

    # Sample Request 5: Document Upload in Postman
    add_code_block(doc,
"""POST {{base_url}}/api/upload
Headers:
  (Leave Content-Type blank - Postman sets multipart/form-data boundary automatically)

Body (form-data):
  file                 -> (Select File: executive_summary.pdf)
  chunking_strategy    -> semantic
  embedding_strategy   -> gemini
  chunk_size           -> 500
  chunk_overlap        -> 100
  extract_graph        -> true""", caption="Example 5: Multipart File Upload in Postman")

    # -------------------------------------------------------------
    # Multi-Level Evaluation & Corruption Suite
    # -------------------------------------------------------------
    p_ev = doc.add_paragraph()
    style_heading_1(p_ev, "6. Multi-Level Evaluation & Corruption Stress Testing")

    add_body_p(doc, 
        "OmniRAG provides an automated diagnostic lab that quantifies system performance across three levels "
        "and stress-tests context degradation under simulated adversarial conditions:"
    )

    t_eval = doc.add_table(rows=1, cols=3)
    t_eval.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_eval,
        [1.8, 1.4, 3.3],
        ["Metric / Diagnostic", "Score Range", "Definition and Significance"],
        [
            [
                "Prompt Clarity Score",
                "0.0 to 1.0",
                "Quantifies question specificity and linguistic clarity. Penalizes ambiguous, vague, or underspecified queries."
            ],
            [
                "Prompt Injection Risk",
                "0.0 to 1.0",
                "Scans for prompt injection attacks, jailbreak phrases (e.g. 'ignore previous instructions'), and system prompt leaks."
            ],
            [
                "Faithfulness Score",
                "0.0 to 1.0",
                "Measures factual consistency: Computes the proportion of claims in the generated response that are directly grounded in retrieved passages."
            ],
            [
                "Hallucination Rate",
                "0.0 to 1.0",
                "Calculated as 1.0 - Faithfulness. Measures the frequency of ungrounded or invented claims."
            ],
            [
                "Answer Relevancy",
                "0.0 to 1.0",
                "Measures semantic cosine similarity between the user's intent and the response, penalizing off-topic tangents."
            ],
            [
                "Format Adherence",
                "0.0 to 1.0",
                "Validates strict compliance with requested output formats (e.g., valid PEP 8 Markdown code blocks or Hindi Devanagari script)."
            ],
            [
                "Corruption Robustness",
                "0.0 to 1.0",
                "Measures the system's ability to maintain accuracy when context passages are corrupted with OCR typos, token dropping, and distractor noise."
            ]
        ]
    )

    # -------------------------------------------------------------
    # Performance & Optimization Suite
    # -------------------------------------------------------------
    p_opt = doc.add_paragraph()
    style_heading_1(p_opt, "7. Latency Profiling & Performance Optimizations")

    add_body_p(doc, 
        "In production enterprise deployments, response latency and infrastructure costs are paramount. "
        "OmniRAG integrates four built-in performance optimization techniques:"
    )

    add_bullet(doc, "1. In-Memory Semantic Vector Response Cache: ", "Maintains an in-memory vector cache of query embeddings and verified responses. Incoming queries with cosine similarity >= 0.92 return in <10ms without invoking Qdrant search or calling external LLM APIs.")
    add_bullet(doc, "2. Context Compression & Sentence Pruning: ", "Applies sentence-level relevance ranking to extract only the most informative sentences from retrieved chunks, reducing prompt token counts by 30-50% and decreasing LLM generation latency.")
    add_bullet(doc, "3. HyDE (Hypothetical Document Embeddings): ", "Synthesizes an ideal hypothetical answer to search for semantically relevant text when the user query is too brief or abstract.")
    add_bullet(doc, "4. Real-Time Telemetry Profiling: ", "Measures fine-grained timing breakdowns for every request: Dense retrieval time (ms), Sparse BM25 search time (ms), RRF fusion time (ms), Time to First Token (TTFT ms), and Tokens Per Second (TPS).")

    # -------------------------------------------------------------
    # Production Deployment & FAQ
    # -------------------------------------------------------------
    p_faq = doc.add_paragraph()
    style_heading_1(p_faq, "8. Production Deployment & Troubleshooting FAQ")

    style_heading_2(doc.add_paragraph(), "8.1 PostgreSQL vs. SQLite Configuration")
    add_body_p(doc, 
        "By default, OmniRAG boots with SQLite fallback if PostgreSQL is not detected. "
        "For production environments with high concurrency, configure PostgreSQL in your .env file or run the automated setup script:"
    )
    add_code_block(doc, "python setup_postgres.py", caption="Automated PostgreSQL Setup Command")

    style_heading_2(doc.add_paragraph(), "8.2 Frequently Asked Questions (FAQ)")

    add_bullet(doc, "Q: Can I run this system completely offline without internet? ", "Yes. In your .env file, set EMBEDDING_PROVIDER=local and point LLM_BACKEND to a local vLLM or SGLang instance. The system will use deterministic 768-d local hashing embeddings and local Qdrant disk storage.")
    add_bullet(doc, "Q: How do I clear the semantic cache if documents are updated? ", "Send a POST request to /api/cache/clear via Swagger UI, Postman, or click the 'Clear Cache' button on the Telemetry tab in the web dashboard.")
    add_bullet(doc, "Q: What file formats are supported for document upload? ", "OmniRAG natively supports PDF (.pdf), CSV (.csv), Excel spreadsheets (.xlsx, .xls), Microsoft Word (.docx, .doc), and plain text (.txt).")
    add_bullet(doc, "Q: How do I inspect active multi-turn sessions? ", "Navigate to GET /api/sessions in Swagger UI or Postman to view all stored sessions, their message counts, and creation timestamps.")

    # -------------------------------------------------------------
    # Verification & Sign-Off
    # -------------------------------------------------------------
    p_sign = doc.add_paragraph()
    style_heading_1(p_sign, "9. Automated Verification & Quality Sign-Off")

    add_body_p(doc, 
        "The OmniRAG platform includes an automated end-to-end test suite in tests/test_all.py covering all 10 major functional subsystems. "
        "All 10 tests execute and pass cleanly in under 70 seconds:"
    )

    t_tests = doc.add_table(rows=1, cols=3)
    t_tests.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_tests,
        [1.8, 1.2, 3.5],
        ["Test Case ID", "Status", "Subsystem Verified"],
        [
            ["test_01_parsers_all_formats", "PASSED", "Verification of document parsers for PDF, CSV, XLSX, DOCX, and TXT."],
            ["test_02_chunking_strategies", "PASSED", "Validation of Recursive, Semantic Sentence, and Structured Table chunkers."],
            ["test_03_embeddings_and_vector_store", "PASSED", "768-dim embeddings validation and Qdrant storage/search."],
            ["test_04_scoring_and_rrf", "PASSED", "BM25Okapi lexical ranking and Reciprocal Rank Fusion calculation (k=60)."],
            ["test_05_three_rag_types", "PASSED", "Execution of Simple RAG, Hybrid RAG, and Graph RAG pipelines."],
            ["test_06_multilingual_and_code_mode", "PASSED", "Grounded Hindi Devanagari synthesis and PEP 8 Python code generation."],
            ["test_07_evaluations_and_corruption", "PASSED", "Prompt evaluation, Response faithfulness, and simulated Corruption stress test."],
            ["test_08_semantic_cache", "PASSED", "Cache hit detection and sub-10ms response retrieval from memory."],
            ["test_09_api_endpoints", "PASSED", "FastAPI REST API endpoints (/api/health, /api/documents, /api/query, /api/graph, /api/benchmark, /api/sessions)."],
            ["test_10_conversational_rag_with_memory", "PASSED", "Multi-turn conversational session memory, coreference resolution, and persistent message tracking."]
        ]
    )

    # -------------------------------------------------------------
    # Multi-Machine Installation & Developer Onboarding Guide
    # -------------------------------------------------------------
    p_inst = doc.add_paragraph()
    style_heading_1(p_inst, "10. Multi-Machine Installation & Developer Onboarding Guide")

    add_body_p(doc, 
        "This chapter provides complete instructions to set up, configure, run, and code on the OmniRAG platform "
        "when moving to a completely new computer (Windows, macOS, or Linux)."
    )

    style_heading_2(doc.add_paragraph(), "10.1 System Prerequisites")
    t_req = doc.add_table(rows=1, cols=3)
    t_req.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table(
        t_req,
        [1.8, 1.4, 3.3],
        ["Component", "Recommended Version", "Notes / Instructions"],
        [
            ["Operating System", "Windows 10/11, macOS 12+, Ubuntu 20.04+", "Fully cross-platform pure Python & FastAPI codebase."],
            ["Python Runtime", "Python 3.10 to 3.12 (3.11 Recommended)", "Python 3.11 provides the fastest execution and pre-built wheels."],
            ["Git", "Git 2.30+", "Required to clone repository and track versions."],
            ["Google Gemini API Key", "Free tier from aistudio.google.com", "Optional: Can run fully offline using local fallback embedder."],
            ["PostgreSQL", "Version 14+ (Optional)", "Not required: Zero-config SQLite automatically acts as fallback."]
        ]
    )

    style_heading_2(doc.add_paragraph(), "10.2 Five-Minute Quickstart (Terminal Commands)")
    add_code_block(doc, 
"""# 1. Clone the repository
git clone <repo-url>
cd RAG_FINAL_PROJECT

# 2. Create and activate a clean virtual environment
python -m venv venv
# Windows:
.\\venv\\Scripts\\activate
# macOS/Linux: source venv/bin/activate

# 3. Upgrade pip and install all dependencies
python -m pip install --upgrade pip
pip install -r requirements.txt

# 4. Copy environment configuration
# Windows: Copy-Item .env.example .env  |  macOS/Linux: cp .env.example .env
# Open .env and add your GEMINI_API_KEY

# 5. Generate test datasets and index them
python create_samples.py
python index.py --all-samples

# 6. Run automated test suite (verifies 10/10 tests)
python -m unittest tests/test_all.py

# 7. Start application server
python main.py""", caption="Standard Multi-Platform Terminal Setup Script")

    style_heading_2(doc.add_paragraph(), "10.3 Developer Blueprint: How to Code & Extend OmniRAG")
    add_body_p(doc, "When extending or developing new capabilities on another computer:")
    add_bullet(doc, "Adding Chunkers (rag/chunking.py): ", "Subclass BaseChunker and register inside get_chunker(strategy_name).")
    add_bullet(doc, "Adding LLM Engines (llm/engine_factory.py): ", "Implement BaseLLMEngine (e.g. Ollama, Claude, DeepSeek) and register in get_llm_engine().")
    add_bullet(doc, "Tuning Scoring Algorithms (rag/scoring.py): ", "Fine-tune BM25 k1 and b values or adjust Reciprocal Rank Fusion k constant.")
    add_bullet(doc, "Writing New Automated Tests (tests/test_all.py): ", "Add new test methods prefixed with test_ and run via python -m unittest.")
    add_bullet(doc, "Modifying the Dashboard (main.py): ", "Frontend HTML, CSS glassmorphic tokens, and Vanilla JS reside inside index_page().")

    style_heading_2(doc.add_paragraph(), "10.4 Troubleshooting Common Platform Gotchas")
    add_bullet(doc, "Windows PowerShell Script Execution: ", "If script activation fails, execute: Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass")
    add_bullet(doc, "Port 8000 Conflict: ", "Run on port 8080 via: uvicorn main:app --port 8080 --reload")
    add_bullet(doc, "Qdrant Directory Lock: ", "Only one process can access ./data/qdrant_storage concurrently. Ensure duplicate background scripts are closed.")
    add_bullet(doc, "PostgreSQL Offline Warning: ", "OmniRAG gracefully falls back to ./data/rag_app.db without crashing. To use PostgreSQL, run python setup_postgres.py.")

    # Save document
    output_path = r"c:\Users\deepa\OneDrive\Desktop\AI ML DATASET\RAG_FINAL_PROJECT\OmniRAG_Complete_User_and_API_Guide.docx"
    doc.save(output_path)
    print(f"Successfully generated documentation file at: {output_path}")

if __name__ == "__main__":
    build_documentation_file()
